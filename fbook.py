import os
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List
import openpyxl
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from langchain.schema import Document  # Import the Document schema

app = FastAPI()

class Question(BaseModel):
    question: str

# Hard code the API key
GOOGLE_API_KEY = "AIzaSyAtkAbHM189Q566Ezh4MyDpdGmEO1Gw_pc"
genai.api_key = GOOGLE_API_KEY

# Global variables to store text chunks and their embeddings
text_chunks = []
embeddings = []

def get_pdf_text(path):
    text = ""
    pr = PdfReader(path)
    for pg in pr.pages:
        text += pg.extract_text()
    return text

def get_docx_text(path):
    text = ""
    doc = docx.Document(path)
    for para in doc.paragraphs:
        text += para.text
    return text

def get_excel_text(path):
    text = ""
    wb = openpyxl.load_workbook(path)
    for sheet in wb:
        for row in sheet.iter_rows(values_only=True):
            text += '\n'.join(map(str, row))
    return text

def get_text_chunks(text):
    ts = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = ts.split_text(text)
    return chunks

def get_embeddings(text_chunks):
    emb = GoogleGenerativeAIEmbeddings(model='models/embedding-001', google_api_key=GOOGLE_API_KEY)
    embeddings = emb.embed_documents(text_chunks)  # Use embed_documents method
    return embeddings

def similarity_search(question, text_chunks, embeddings):
    emb = GoogleGenerativeAIEmbeddings(model='models/embedding-001', google_api_key=GOOGLE_API_KEY)
    question_embedding = emb.embed_query(question)  # Use embed_query method for single query

    similarities = cosine_similarity([question_embedding], embeddings)
    most_similar_idx = np.argmax(similarities)

    return text_chunks[most_similar_idx]

def get_conversational_chain():
    prmt_temp = """
    Answer the question as detailed and as thoroughly as possible using the provided context. Never answer with just the answer always frame sentences. If asked to summarize shorten the entire thing and give answer.
    If the answer is not found in the context, say "Unable to find the answer", but do not provide incorrect information.
    Be comprehensive in your response and provide as much detail as possible.
    context:\n{context}\n
    Question:\n{question}\n
    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.7, google_api_key=GOOGLE_API_KEY)
    prompt = PromptTemplate(template=prmt_temp, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

def user_input(ques):
    most_relevant_chunk = similarity_search(ques, text_chunks, embeddings)
    # Wrap the text chunk in a Document object
    doc = Document(page_content=most_relevant_chunk, metadata={})
    chain = get_conversational_chain()
    res = chain.invoke(
        {"input_documents": [doc], "question": ques},
    )
    return res["output_text"]

@app.post("/upload/")
async def upload_file(files: List[UploadFile] = File(...)):
    global text_chunks, embeddings
    mas_read = ""
    for file in files:
        if file.filename.endswith('.pdf'):
            mas_read += get_pdf_text(file.file)
        elif file.filename.endswith('.docx'):
            mas_read += get_docx_text(file.file)
        elif file.filename.endswith('.xlsx'):
            mas_read += get_excel_text(file.file)

    text_chunks = get_text_chunks(mas_read)
    embeddings = get_embeddings(text_chunks)
    return {"filenames": [file.filename for file in files]}

@app.post("/ask/")
async def ask_question(question: Question):
    answer = user_input(question.question)
    return JSONResponse(content={"question": question.question, "answer": answer})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
